from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from fastapi.responses import FileResponse
from reportlab.pdfgen import canvas
from pathlib import Path
from typing import Optional
from uuid import uuid4
from datetime import datetime
import os
import base64
import binascii
import mimetypes

# For extracting content
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader

app = FastAPI()

# ------------------------------------------------------------------------------
# Configuration
# ------------------------------------------------------------------------------

BASE_URL = os.getenv(
    "BASE_URL",
    "https://python-document-new-2.onrender.com"
)

FILES_DIR = Path(os.getcwd()) / "generated_files"
FILES_DIR.mkdir(parents=True, exist_ok=True)

# ------------------------------------------------------------------------------
# Request model
# ------------------------------------------------------------------------------

class GenerateRequest(BaseModel):
    project_id: int
    project_plan_line_id: int
    template_id: str
    template_content: str
    template_file_name: Optional[str] = None


# ------------------------------------------------------------------------------
# Content extraction helper
# ------------------------------------------------------------------------------

def extract_file_content(file_path: Path) -> str:
    """
    Extract readable content from supported file types.
    """

    suffix = file_path.suffix.lower()

    # TXT / JSON / XML / HTML / CSV
    if suffix in [".txt", ".json", ".xml", ".html", ".csv"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    # DOCX
    elif suffix == ".docx":
        doc = Document(str(file_path))
        lines = []

        # paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())

        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    lines.append(" | ".join(row_text))

        return "\n".join(lines)

    # XLSX
    elif suffix == ".xlsx":
        wb = load_workbook(str(file_path), data_only=True)
        all_text = []

        for sheet in wb.worksheets:
            all_text.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    all_text.append(" | ".join(row_values))

        return "\n".join(all_text)

    # PDF
    elif suffix == ".pdf":
        reader = PdfReader(str(file_path))
        pages_text = []

        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages_text.append(f"--- Page {i} ---\n{text}")

        return "\n".join(pages_text)

    # Unsupported / binary
    else:
        return "Unsupported file type for direct text extraction."


# ------------------------------------------------------------------------------
# Generate endpoint
# ------------------------------------------------------------------------------

@app.post("/api/v1/deliverables/generate")
def generate(req: GenerateRequest):
    """
    Receives base64 template content from APEX,
    saves original file,
    extracts readable content,
    creates preview PDF,
    returns extracted content + links.
    """

    try:
        # Step 1: decode base64
        decoded_bytes = base64.b64decode(req.template_content)

    except binascii.Error:
        return {
            "status": "ERROR",
            "message": "Invalid base64 template content received from APEX"
        }

    try:
        # Step 2: create unique names for every run
        unique_id = f"{req.project_plan_line_id}_{req.template_id}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}"

        input_name = req.template_file_name or f"template_{unique_id}.bin"
        input_name = Path(input_name).name
        suffix = Path(input_name).suffix or ".bin"

        original_file_name = f"template_{unique_id}{suffix}"
        original_file_path = FILES_DIR / original_file_name

        with open(original_file_path, "wb") as f:
            f.write(decoded_bytes)

        # Step 3: extract inside content
        extracted_content = extract_file_content(original_file_path)

        # Step 4: save extracted text as .txt file
        text_file_name = f"content_{unique_id}.txt"
        text_file_path = FILES_DIR / text_file_name

        with open(text_file_path, "w", encoding="utf-8") as f:
            f.write(extracted_content)

        # Step 5: create preview PDF with extracted text
        preview_file_name = f"preview_{unique_id}.pdf"
        preview_file_path = FILES_DIR / preview_file_name

        c = canvas.Canvas(str(preview_file_path))
        c.setTitle("Template Preview")

        y = 800
        c.drawString(50, y, "Template Extracted Content Preview")
        y -= 30
        c.drawString(50, y, f"Project ID: {req.project_id}")
        y -= 20
        c.drawString(50, y, f"Project Plan Line ID: {req.project_plan_line_id}")
        y -= 20
        c.drawString(50, y, f"Template ID: {req.template_id}")
        y -= 20
        c.drawString(50, y, f"Original File Name: {original_file_name}")
        y -= 30

        lines = extracted_content.splitlines()
        for line in lines[:40]:
            c.drawString(50, y, line[:110])
            y -= 18
            if y < 50:
                c.showPage()
                y = 800

        c.save()

        return {
            "status": "SUCCESS",
            "message": "Template content extracted successfully",
            "type": "PDF",
            "file_name": preview_file_name,
            "link": f"{BASE_URL}/api/v1/deliverables/download/{preview_file_name}",
            "original_file_name": original_file_name,
            "original_file_link": f"{BASE_URL}/api/v1/deliverables/download/{original_file_name}",
            "text_file_name": text_file_name,
            "text_file_link": f"{BASE_URL}/api/v1/deliverables/download/{text_file_name}",
            "extracted_content": extracted_content[:4000]
        }

    except Exception as e:
        return {
            "status": "ERROR",
            "message": str(e)
        }


# ------------------------------------------------------------------------------
# Download endpoint
# ------------------------------------------------------------------------------

@app.get("/api/v1/deliverables/download/{file_name}")
def download(file_name: str):
    safe_file_name = Path(file_name).name
    file_path = FILES_DIR / safe_file_name

    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    media_type, _ = mimetypes.guess_type(str(file_path))
    if not media_type:
        media_type = "application/octet-stream"

    return FileResponse(
        path=str(file_path),
        media_type=media_type,
        filename=safe_file_name,
        headers={
            "Content-Disposition": f'attachment; filename="{safe_file_name}"',
            "Cache-Control": "no-store"
        }
    )
